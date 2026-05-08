from docx import Document as DocxDocument

from app.schemas.parser import (
    ParsedDocument,
    ParsedPage,
    ParsedSection,
    ParsedTable,
)


class DOCXParser:

    @staticmethod
    def parse(
        file_path: str,
    ) -> ParsedDocument:

        try:

            document = DocxDocument(file_path)

            parsed_pages = []
            parsed_sections = []
            parsed_tables = []

            extracted_content = []

            current_section_title = "Introduction"
            current_section_content = []

            # Parse paragraphs
            for paragraph in document.paragraphs:

                text = paragraph.text.strip()

                if not text:
                    continue

                style_name = (
                    paragraph.style.name.lower()
                )

                extracted_content.append(text)

                if "heading" in style_name:

                    if current_section_content:

                        parsed_sections.append(
                            ParsedSection(
                                title=current_section_title,
                                content="\n".join(
                                    current_section_content
                                ),
                                page_number=1,
                            )
                        )

                    current_section_title = text

                    current_section_content = []

                else:
                    current_section_content.append(
                        text
                    )

            # Final section
            if current_section_content:

                parsed_sections.append(
                    ParsedSection(
                        title=current_section_title,
                        content="\n".join(
                            current_section_content
                        ),
                        page_number=1,
                    )
                )

            # Parse tables
            for table in document.tables:

                headers = []
                rows = []

                for row_index, row in enumerate(
                    table.rows
                ):

                    row_data = []

                    for cell in row.cells:

                        row_data.append(
                            cell.text.strip()
                        )

                    if row_index == 0:
                        headers = row_data
                    else:
                        rows.append(row_data)

                parsed_tables.append(
                    ParsedTable(
                        title=None,
                        headers=headers,
                        rows=rows,
                        page_number=1,
                    )
                )

            parsed_pages.append(
                ParsedPage(
                    page_number=1,
                    text="\n".join(
                        extracted_content
                    ),
                )
            )

            return ParsedDocument(
                sections=parsed_sections,
                tables=parsed_tables,
                pages=parsed_pages,
                total_pages=len(parsed_pages),
            )

        except Exception as exc:
            raise ValueError(
                f"DOCX parsing failed: {exc}"
            )